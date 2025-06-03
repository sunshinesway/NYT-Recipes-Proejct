from bs4 import BeautifulSoup
import json
import os
from PIL import Image 
import requests
	
import docx
# from docx.shared import Inches

class Recipe:
    def __init__(self, ingred, steps, picture, title_info):
        self.ingred = ingred # ingredient list []
        self.steps = steps # prep steps list []
        self.picture = picture # image file path
        self.title_info = title_info # dict: {title, descr, time, yield, category, list_name}

def rem_html_video(filename):
        # remove .html
        update = filename.replace(".html", "", 1)
        # remove (with Video) if present
        update_1 = update.replace("(with Video)", "", 1)

        return update_1

# add recipe items in dir to list
def recipe_list():
    path = "./NYT-recipes"
    dir_list = os.listdir(path)

    recipes = []
    for item in dir_list:
        print(item)
        with open(os.path.join(path, item), 'r') as f:
            if item.find('.html') != -1:             
                html_string = f.read()
            else:
                continue
        
        soup = BeautifulSoup(html_string, 'html.parser')
        recipe_json = soup.script.string
        parse_json = json.loads(recipe_json)

        header_info = {'title': parse_json["name"],
                       'descr': parse_json["description"],
                       'totalTime': parse_json["totalTime"],
                       'yield': parse_json["recipeYield"],
                       'category': parse_json["recipeCategory"],
                       'list_name': rem_html_video(item) }

        ingred_list = parse_json["recipeIngredient"]
        instructions = parse_json["recipeInstructions"]

        steps = []
        for step in instructions:
            steps.append(step["text"])

        image_json = parse_json["image"]
        picture_link = image_json[2]["contentUrl"]
        img_items = [item, '-image.png']
        img_file = ''.join(img_items)
        # urllib.request.urlretrieve(picture_link, img_file)
        data = requests.get(picture_link).content
        with open(f'./images/{header_info['list_name']}-image.jpg', 'wb') as f:
            f.write(data)
            f.close

        new_recipe = Recipe(ingred_list, steps, data, header_info)
        recipes.append(new_recipe)
        print(f"{new_recipe.title_info['title']} added to list")

    return recipes


# recipes = recipe_list()

# for item in recipes:
#     # Create a document
#     doc = docx.Document()

#     # Add a paragraph to the document
#     p = doc.add_paragraph()

#     # Add some formatting to the paragraph
#     p.paragraph_format.line_spacing = 1
#     p.paragraph_format.space_after = 0

#     # Add a run to the paragraph for title
#     run = p.add_run(item.title_info['title'])

#     # Add some formatting to the run
#     run.bold = True
#     run.font.name = 'Arial'
#     run.font.size = docx.shared.Pt(20)

#     # empty line
#     doc.add_paragraph()

#     # Add another paragraph
#     p = doc.add_paragraph()

#     # Add a run Ingredients
#     run = p.add_run("INGREDIENTS\n")
#     run.font.name = 'Arial'
#     run.bold = True
#     run.font.size = docx.shared.Pt(16)

#     # Add a run Yield
#     yield_str = f"Yield: {item.title_info['yield']}\n"
#     run = p.add_run(yield_str)
#     run.font.name = 'Arial'
#     # run.bold = True
#     run.font.size = docx.shared.Pt(14)

#     for ingredient in item.ingred:
#         # Add a run for each ingredient
#         ingredient_str = f"    - {ingredient}\n"
#         run = p.add_run(ingredient_str)
#         run.font.name = 'Arial'
#         run.font.size = docx.shared.Pt(14)
    
#     doc.add_paragraph() # empty line

#     # Add another paragraph
#     p = doc.add_paragraph()

#     # Add a run Ingredients
#     run = p.add_run("PREPARATION\n")
#     run.font.name = 'Arial'
#     run.bold = True
#     run.font.size = docx.shared.Pt(16)
#     run.paragraph_format.line_spacing

#     index = 1
#     for step in item.steps:
#         # Add a run for each step
#         step_str = f"{index}. {step}\n"
#         run = p.add_run(step_str)
#         run.font.name = 'Arial'
#         run.font.size = docx.shared.Pt(14)
#         index += 1

#     # Save the document
#     doc.save("./docs/", f"{item.title_info['list_name']}.docx")

# with open('Mexican-Rice-Recipe-only.html', 'r') as f:
#     html_string = f.read()    

# soup = BeautifulSoup(html_string, 'html.parser')

# recipe_json = soup.script.string

# parse_json = json.loads(recipe_json)

# ingre_list = parse_json["recipeIngredient"]
# for item in ingre_list:
#     print(item)

# instructions = parse_json["recipeInstructions"]

# for step in instructions:
#     print(step["text"])

# print("\n\n")
# # list files in directory
# path = "./NYT-recipes"
# dir_list = os.listdir(path)

# # create new list for recipe names - .html
# recipe_list = []
# for item in dir_list:
#     # remove .html
#     update = item.replace(".html", "", 1)
#     # remove (with Video) if present
#     if update.find("(with Video)") != -1 :
#         update_1 = update.replace("(with Video)", "", 1)
#         update = update_1
#     # add to list
#     recipe_list.append(update)

# print(f"Files in '{path}' : ")
# print(recipe_list)

# img = Image.open(f"{recipes[0].title_info['list_name']}-image.png")



