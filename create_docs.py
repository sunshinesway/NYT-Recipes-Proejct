import docx
from docx.shared import Inches
from docx.shared import Pt

def docx_template(recipe):

    # Create document
    doc = docx.Document()

    # set margins
    # leave default 1" top & bottom
    sections = doc.sections
    for section in sections:
        section.right_margin = Inches(1)
        section.left_margin = Inches(1)

    # Add heading
    doc.add_heading(f'{recipe.title_info['title']}', 0)

    # add yield
    doc.add_heading(f'Yield: {recipe.title_info['yield']}', 4)

    # add Ingredients heading
    doc.add_heading('INGREDIENTS', 3)

    # add ingredient list
    for item in recipe.ingred:
        p = doc.add_paragraph(f'   - {item}')
        p.paragraph_format.space_after = Inches(0.1)
        p.style.font.name = 'Calibri'
        p.style.font.size = Pt(13)
        # p.font.name = 'Calibri'
    
    # add Preparation heading
    doc.add_heading('PREPARATION', 3)

    # add prep steps list
    index = 1
    for item in recipe.steps:
        p = doc.add_paragraph(f'{index}.   {item}')
        p.paragraph_format.space_after = Inches(0.1)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.keep_together = True
        p.style.font.name = 'Calibri'
        p.style.font.size = Pt(13)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        index += 1
    
    doc.save(f'./docs/{recipe.title_info['list_name']}.docx')



